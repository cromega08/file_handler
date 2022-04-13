import PyPDF4 as pdf_r
import docx as word
import fpdf as pdf_w
import argparse as argp
import os
class files():

    def __init__(self, filename):
        
        self.filename = filename
        self.root, self.ext = os.path.splitext(self.filename)

    def read_pdf(self):

        pdf_file = os.open(self.filename, os.O_RDWR)

        pdf_reader = pdf_r.PdfFileReader()
        
        if pdf_reader.isEncrypted: pdf_reader.decrypt(self.get_pdf_pswd())

        pure_text = []

        for pages in range(pdf_reader.numPages): pure_text.append(pdf_reader.getPage(pages).extractText())
            #caution: if output error, look the example in book
        
        return pure_text

    def get_pdf_pswd(self):

        pswd = input("It's seem the file is encrypted\n\nPassword of the .pdf file: ")

        return pswd

    def write_pdf(self):

        pass

    def create_pdf_pages(self):

        new_pdf = pdf_w()
        new_pdf.add_page()
        

    def read_word(self):

        word_reader = word.Document(self.filename)
        full_text = []

        for paragraphs in word_reader.paragraphs: full_text.append(paragraphs), full_text.append("\n")

        return full_text

    def write_word(self, full_text):

        new_word = word.Document()

        for text in full_text: new_word.add_paragraph(text)

        new_word.save(f"{self.root}.docx")

class cli():

    my_parser = argp.ArgumentParser(prog = "fyle", description = "A path.exist CLI", usage = "%(prog)s [options] path")
    my_parser.add_argument("path", type = str, help = "the path to confirm his existence")
    arg = my_parser.parse_args()

    exist = os.path.exists(arg.path)
